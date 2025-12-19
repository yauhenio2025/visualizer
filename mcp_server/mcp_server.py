"""
Visualizer MCP Server

Exposes document analysis and visualization capabilities via Model Context Protocol.

Usage:
    claude mcp add --transport stdio visualizer -- python /path/to/mcp_server.py

Configuration via environment variables:
    VISUALIZER_API_URL: Visualizer API base URL (default: https://visualizer-tw4i.onrender.com)
    ANALYZER_API_URL: Analyzer API base URL (default: https://analyzer-3wsg.onrender.com)
    GEMINI_API_KEY: User's Gemini API key (optional, can be provided per-request)
    ANTHROPIC_API_KEY: User's Anthropic API key (optional, can be provided per-request)
    VISUALIZER_NTFY_TOPIC: Notification topic for job completion alerts
    VISUALIZER_OUTPUT_DIR: Directory for downloaded results (default: ~/visualizer-results)

Primary Usage Pattern:
    1. User provides a document path or content
    2. Claude Code calls get_ai_recommendations() to suggest best engines
    3. User selects which engines to use (or accepts all)
    4. Claude Code calls submit_analysis() with selected engines
    5. Server polls job status and sends notification when complete
    6. Results are downloaded from S3 and saved locally
    7. User is alerted with sound signal
"""

import os
import sys
import time
import json
import logging
import requests
import base64
import subprocess
from datetime import datetime
from typing import Annotated, Optional, List, Dict, Any
from pathlib import Path

# Load .env.mcp from the same directory as this script BEFORE other imports
# This ensures API keys are available when the MCP server starts
_script_dir = Path(__file__).parent
_env_file = _script_dir / ".env.mcp"
if _env_file.exists():
    with open(_env_file) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                key, value = line.split('=', 1)
                os.environ.setdefault(key.strip(), value.strip())

from fastmcp import FastMCP
from pydantic import Field

# PDF support
try:
    import fitz  # pymupdf
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logging.warning("pymupdf not installed - PDF text extraction disabled")

# DOCX support
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logging.warning("python-docx not installed - DOCX text extraction disabled")

# Configure logging to stderr (stdout is reserved for JSON-RPC) AND file
LOG_LEVEL = os.environ.get('VISUALIZER_LOG_LEVEL', 'DEBUG').upper()
LOG_FILE = os.environ.get('VISUALIZER_LOG_FILE', str(Path.home() / '.visualizer-mcp.log'))

# Create formatter with detailed information
log_formatter = logging.Formatter(
    '%(asctime)s | %(levelname)-8s | %(funcName)-25s | %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# Configure root logger
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL, logging.DEBUG),
    format='%(asctime)s | %(levelname)-8s | %(funcName)-25s | %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
    handlers=[logging.StreamHandler(sys.stderr)]
)

# Add file handler for persistent logs with immediate flush
class FlushingFileHandler(logging.FileHandler):
    """File handler that flushes after every log message - critical for MCP subprocess context."""
    def emit(self, record):
        super().emit(record)
        self.flush()

try:
    file_handler = FlushingFileHandler(LOG_FILE, mode='a', encoding='utf-8')
    file_handler.setFormatter(log_formatter)
    file_handler.setLevel(logging.DEBUG)  # Always log everything to file
    logging.getLogger().addHandler(file_handler)
    # Also add to all loggers that might be used
    logging.getLogger('__main__').addHandler(file_handler)
except Exception as e:
    sys.stderr.write(f"Warning: Could not create log file {LOG_FILE}: {e}\n")

logger = logging.getLogger(__name__)
logger.info(f"=== MCP Server Starting === Log level: {LOG_LEVEL}, Log file: {LOG_FILE}")
logger.info(f"Python: {sys.version}, PID: {os.getpid()}")

# Initialize MCP server
mcp = FastMCP(
    name="visualizer",
    instructions="""
    Visualizer MCP Server - Document Intelligence & Visual Analysis

    This server connects Claude Code to a powerful document analysis and visualization
    pipeline that uses AI to recommend the best analysis engines for your documents.

    ## Primary Usage

    When a user wants to analyze/visualize a document:

    1. **Get AI Recommendations**: Call get_ai_recommendations() with document path/content
       - Returns 4-5 recommended engines with confidence scores
    2. **Submit Analysis**: Call submit_analysis() with selected engines
       - Supports both visual and textual output modes
    3. **Monitor Progress**: Call check_job_status() to track progress
    4. **Get Results**: Results are auto-downloaded when complete, or use get_results()

    ## Features

    - AI-powered engine recommendations using Claude/Gemini
    - 70 analysis engines across 9 categories (Argument, Concepts, Temporal, Power, Evidence, Rhetoric, Epistemology, Scholarly, Market)
    - 18 engine bundles for thematic analysis
    - 21 multi-stage pipelines that chain engines sequentially
    - Multiple output modes: visual (4K images) and textual (reports, diagrams)
    - Batch analysis: Process entire folders of documents
    - Automatic notifications with sound alerts when jobs complete
    """
)

# Configuration from environment
VISUALIZER_API_URL = os.environ.get('VISUALIZER_API_URL', 'https://visualizer-tw4i.onrender.com')
ANALYZER_API_URL = os.environ.get('ANALYZER_API_URL', 'https://analyzer-3wsg.onrender.com')
OUTPUT_DIR = Path(os.environ.get('VISUALIZER_OUTPUT_DIR', '~/visualizer-results')).expanduser()
NTFY_TOPIC = os.environ.get('VISUALIZER_NTFY_TOPIC', f'visualizer-{os.getenv("USER", "user")}-{hash(os.getenv("USER", "user")) % 10000}')

# Ensure output directory exists
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Request timeout
REQUEST_TIMEOUT = 30
POLL_INTERVAL = 5  # seconds
MAX_POLL_ATTEMPTS = 600  # 30 minutes max

# Supported file extensions for batch processing
SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.md', '.markdown', '.rst', '.tex', '.docx'}


def sanitize_api_key(key: Optional[str]) -> Optional[str]:
    """Strip whitespace/newlines from API keys that may come from display formatting."""
    if key:
        # Remove all whitespace including newlines, spaces, tabs
        return ''.join(key.split())
    return key


def get_llm_keys(anthropic_api_key: Optional[str] = None, gemini_api_key: Optional[str] = None) -> Dict[str, str]:
    """Build llm_keys dict from provided keys or environment."""
    llm_keys = {}

    # Sanitize keys to remove any whitespace/newlines from display formatting
    anthropic_key = sanitize_api_key(anthropic_api_key) or sanitize_api_key(os.environ.get('ANTHROPIC_API_KEY'))
    gemini_key = sanitize_api_key(gemini_api_key) or sanitize_api_key(os.environ.get('GEMINI_API_KEY'))

    if anthropic_key:
        llm_keys['anthropic_api_key'] = anthropic_key
    if gemini_key:
        llm_keys['gemini_api_key'] = gemini_key

    return llm_keys


def api_request(
    base_url: str,
    method: str,
    endpoint: str,
    data: dict = None,
    headers: dict = None,
    timeout: int = REQUEST_TIMEOUT
) -> dict:
    """Make an API request to visualizer or analyzer backend."""
    url = f"{base_url}{endpoint}"
    request_id = f"{method}_{endpoint}_{int(time.time() * 1000) % 100000}"

    # Log request details (sanitize sensitive data)
    safe_data = None
    if data:
        safe_data = {k: ('***' if 'key' in k.lower() or 'token' in k.lower() else
                        f"<{len(str(v))} chars>" if isinstance(v, str) and len(str(v)) > 200 else v)
                     for k, v in data.items()}

    logger.debug(f"[{request_id}] >>> API Request: {method} {url}")
    logger.debug(f"[{request_id}] >>> Data: {json.dumps(safe_data, default=str)[:500] if safe_data else 'None'}")
    logger.debug(f"[{request_id}] >>> Timeout: {timeout}s")

    start_time = time.time()

    try:
        request_headers = headers or {}

        if method.upper() == 'GET':
            response = requests.get(url, params=data, headers=request_headers, timeout=timeout)
        elif method.upper() == 'POST':
            response = requests.post(url, json=data, headers=request_headers, timeout=timeout)
        elif method.upper() == 'PUT':
            response = requests.put(url, json=data, headers=request_headers, timeout=timeout)
        elif method.upper() == 'DELETE':
            response = requests.delete(url, headers=request_headers, timeout=timeout)
        else:
            logger.error(f"[{request_id}] Unsupported HTTP method: {method}")
            return {"error": f"Unsupported HTTP method: {method}"}

        elapsed = time.time() - start_time
        logger.debug(f"[{request_id}] <<< Response: {response.status_code} in {elapsed:.2f}s")

        response.raise_for_status()
        result = response.json()

        # Log response summary
        if isinstance(result, dict):
            if 'error' in result:
                logger.warning(f"[{request_id}] <<< API returned error: {result.get('error')}")
            elif 'job_id' in result:
                logger.info(f"[{request_id}] <<< Job ID: {result.get('job_id')}, Status: {result.get('status', 'N/A')}")
            elif 'status' in result:
                logger.debug(f"[{request_id}] <<< Status: {result.get('status')}")

        return result

    except requests.Timeout:
        elapsed = time.time() - start_time
        logger.error(f"[{request_id}] !!! TIMEOUT after {elapsed:.2f}s (limit: {timeout}s) - URL: {url}")
        return {"error": f"Request timeout after {timeout}s", "url": url, "request_id": request_id}

    except requests.HTTPError as e:
        elapsed = time.time() - start_time
        error_text = ""
        status_code = getattr(e.response, 'status_code', 'unknown')
        try:
            error_text = e.response.text[:1000]
        except:
            pass
        logger.error(f"[{request_id}] !!! HTTP {status_code} after {elapsed:.2f}s - URL: {url}")
        logger.error(f"[{request_id}] !!! Response body: {error_text[:500]}")
        return {"error": f"API error: {status_code} {str(e)}", "details": error_text, "url": url, "request_id": request_id}

    except requests.RequestException as e:
        elapsed = time.time() - start_time
        logger.error(f"[{request_id}] !!! Request failed after {elapsed:.2f}s: {type(e).__name__}: {str(e)}")
        logger.exception(f"[{request_id}] Full exception:")
        return {"error": f"Request failed: {str(e)}", "url": url, "request_id": request_id}

    except Exception as e:
        elapsed = time.time() - start_time
        logger.error(f"[{request_id}] !!! Unexpected error after {elapsed:.2f}s: {type(e).__name__}: {str(e)}")
        logger.exception(f"[{request_id}] Full exception:")
        return {"error": f"Unexpected error: {str(e)}", "url": url, "request_id": request_id}


def extract_pdf_text(file_path: Path) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    logger.debug(f"Extracting text from PDF: {file_path}")

    if not PDF_SUPPORT:
        logger.warning(f"PDF extraction not available (pymupdf not installed)")
        return "[PDF text extraction not available - install pymupdf]"

    try:
        doc = fitz.open(str(file_path))
        page_count = len(doc)
        logger.debug(f"PDF opened: {page_count} pages")

        text_parts = []
        chars_extracted = 0
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                chars_extracted += len(text)
        doc.close()

        if text_parts:
            logger.info(f"PDF extracted: {len(text_parts)}/{page_count} pages, {chars_extracted:,} chars")
            return "\n\n".join(text_parts)
        else:
            logger.warning(f"PDF has NO extractable text (may be scanned images): {file_path}")
            return "[No text extracted from PDF - may be scanned images without OCR]"

    except Exception as e:
        logger.error(f"PDF extraction failed: {type(e).__name__}: {str(e)}")
        logger.exception("Full PDF extraction error:")
        return f"[Error extracting PDF text: {str(e)}]"


def extract_docx_text(file_path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    logger.debug(f"Extracting text from DOCX: {file_path}")

    if not DOCX_SUPPORT:
        logger.warning(f"DOCX extraction not available (python-docx not installed)")
        return "[DOCX text extraction not available - install python-docx]"

    try:
        doc = DocxDocument(str(file_path))
        text_parts = []
        para_count = 0
        table_count = 0

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                para_count += 1

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    table_count += 1

        if text_parts:
            total_chars = sum(len(t) for t in text_parts)
            logger.info(f"DOCX extracted: {para_count} paragraphs, {table_count} table rows, {total_chars:,} chars")
            return "\n\n".join(text_parts)
        else:
            logger.warning(f"DOCX has no extractable text: {file_path}")
            return "[No text extracted from DOCX]"

    except Exception as e:
        logger.error(f"DOCX extraction failed: {type(e).__name__}: {str(e)}")
        logger.exception("Full DOCX extraction error:")
        return f"[Error extracting DOCX text: {str(e)}]"


def read_document(file_path: str) -> Dict[str, Any]:
    """Read a document from file path and return structured document object."""
    logger.info(f"Reading document: {file_path}")
    path = Path(file_path).expanduser().resolve()

    if not path.exists():
        logger.error(f"File not found: {file_path} (resolved: {path})")
        return {"error": f"File not found: {file_path}"}

    if not path.is_file():
        logger.error(f"Not a file: {file_path}")
        return {"error": f"Not a file: {file_path}"}

    file_size = path.stat().st_size
    logger.debug(f"File exists: {path}, size: {file_size:,} bytes")

    # Determine file type
    suffix = path.suffix.lower()
    logger.debug(f"File type: {suffix}")

    # For PDFs, read as base64 AND extract text
    if suffix == '.pdf':
        try:
            logger.debug(f"Reading PDF as base64...")
            with open(path, 'rb') as f:
                content = base64.b64encode(f.read()).decode('utf-8')
            encoding = 'base64'
            logger.debug(f"PDF base64 size: {len(content):,} chars")
            # Also extract text for sample_text usage
            extracted_text = extract_pdf_text(path)
            logger.debug(f"Extracted text length: {len(extracted_text):,} chars")
        except Exception as e:
            logger.error(f"Failed to read PDF: {type(e).__name__}: {str(e)}")
            logger.exception("Full PDF read error:")
            return {"error": f"Cannot read PDF: {str(e)}"}
    # For DOCX, read as base64 AND extract text
    elif suffix == '.docx':
        try:
            logger.debug(f"Reading DOCX as base64...")
            with open(path, 'rb') as f:
                content = base64.b64encode(f.read()).decode('utf-8')
            encoding = 'base64'
            logger.debug(f"DOCX base64 size: {len(content):,} chars")
            # Also extract text for sample_text usage
            extracted_text = extract_docx_text(path)
            logger.debug(f"Extracted text length: {len(extracted_text):,} chars")
        except Exception as e:
            logger.error(f"Failed to read DOCX: {type(e).__name__}: {str(e)}")
            logger.exception("Full DOCX read error:")
            return {"error": f"Cannot read DOCX: {str(e)}"}
    else:
        # For text files (txt, md, markdown, rst, tex), read as text
        try:
            logger.debug(f"Reading text file...")
            content = path.read_text(encoding='utf-8')
            encoding = 'text'
            extracted_text = content  # Same as content for text files
            logger.debug(f"Text content length: {len(content):,} chars")
        except UnicodeDecodeError as e:
            logger.error(f"Unicode decode error for {path}: {str(e)}")
            return {"error": f"Cannot read file as text (not a supported format): {file_path}"}
        except Exception as e:
            logger.error(f"Failed to read text file: {type(e).__name__}: {str(e)}")
            logger.exception("Full text read error:")
            return {"error": f"Cannot read file: {str(e)}"}

    doc_id = f"doc_{int(time.time())}_{hash(str(path)) % 10000}"
    logger.info(f"Document ready: {doc_id}, title: {path.name}, size: {file_size:,} bytes, encoding: {encoding}")

    return {
        "id": doc_id,
        "title": path.name,
        "content": content,
        "encoding": encoding,
        "extracted_text": extracted_text,  # Always have text available
        "path": str(path),
        "size": file_size
    }


def send_notification(title: str, message: str, tags: str = "visualizer", sound: bool = True):
    """Send notification via ntfy.sh.

    Args:
        title: Notification title
        message: Notification body
        tags: Comma-separated tags
        sound: If True, plays sound (priority 4). If False, silent (priority 3).
    """
    try:
        requests.post(
            f"https://ntfy.sh/{NTFY_TOPIC}",
            json={
                "topic": NTFY_TOPIC,
                "title": title,
                "message": message,
                "tags": tags,
                "priority": 4 if sound else 1  # 4=high (sound), 1=min (truly silent)
            },
            timeout=5
        )
    except Exception as e:
        logger.warning(f"Failed to send notification: {e}")


def spawn_job_poller(job_ids: List[str]):
    """Spawn job_poller.py as a background process to monitor and auto-download results.

    Args:
        job_ids: List of job IDs to monitor
    """
    if not job_ids:
        return

    # Get path to job_poller.py (in same directory as this script)
    poller_path = Path(__file__).parent / "job_poller.py"

    if not poller_path.exists():
        logger.warning(f"job_poller.py not found at {poller_path}")
        return

    try:
        # Spawn poller as detached background process
        # Uses nohup + disown pattern to fully detach
        cmd = [sys.executable, str(poller_path)] + job_ids

        # Open /dev/null for stdin, redirect stdout/stderr to log file
        log_file = OUTPUT_DIR / "poller.log"

        with open(log_file, 'a') as log:
            log.write(f"\n{'='*50}\n")
            log.write(f"[{datetime.now().isoformat()}] Starting poller for {len(job_ids)} jobs\n")
            log.write(f"Jobs: {job_ids}\n")
            log.flush()

            process = subprocess.Popen(
                cmd,
                stdout=log,
                stderr=subprocess.STDOUT,
                stdin=subprocess.DEVNULL,
                start_new_session=True  # Detach from parent process group
            )

        logger.info(f"Spawned job_poller (PID {process.pid}) for {len(job_ids)} jobs")
    except Exception as e:
        logger.error(f"Failed to spawn job_poller: {e}")


def download_file_from_url(url: str, output_path: Path) -> tuple[bool, str]:
    """Download a file from URL to the specified path.

    Returns: (success, error_message)
    """
    start_time = time.time()
    # Truncate URL for logging (hide signature params)
    url_display = url.split('?')[0][-60:] if '?' in url else url[-60:]
    logger.debug(f"Downloading: ...{url_display}")

    try:
        response = requests.get(url, timeout=120, stream=True)
        response.raise_for_status()

        content_length = response.headers.get('content-length')
        if content_length:
            logger.debug(f"Content-Length: {int(content_length):,} bytes")

        total_bytes = 0
        with open(output_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
                total_bytes += len(chunk)

        elapsed = time.time() - start_time
        speed_kbps = (total_bytes / 1024) / elapsed if elapsed > 0 else 0
        logger.debug(f"Downloaded {total_bytes:,} bytes in {elapsed:.1f}s ({speed_kbps:.0f} KB/s)")
        return True, ""
    except requests.Timeout:
        elapsed = time.time() - start_time
        error_msg = f"Timeout after {elapsed:.1f}s"
        logger.error(f"Download timeout: {error_msg} - URL: ...{url_display}")
        return False, error_msg
    except requests.HTTPError as e:
        status = getattr(e.response, 'status_code', 'unknown')
        error_msg = f"HTTP {status}: {str(e)[:100]}"
        logger.error(f"Download HTTP error: {error_msg}")
        return False, error_msg
    except Exception as e:
        error_msg = f"{type(e).__name__}: {str(e)[:100]}"
        logger.error(f"Download failed: {error_msg} - URL: ...{url_display}")
        return False, error_msg


def generate_meaningful_folder_name(result: dict, job_id: str) -> str:
    """Generate a meaningful folder name from job result data.

    Format: YYYYMMDD_HHMMSS_engine_DocumentTitle
    Example: 20251218_200113_dialectical_structure_Four_Forms_Critical_Theory

    Args:
        result: Job result dict containing extended_info, outputs, legend
        job_id: Job ID (used as fallback if no meaningful name can be generated)

    Returns:
        A sanitized folder name string
    """
    import re

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Extract engine name(s) from outputs or extended_info
    engine_parts = []
    outputs = result.get("outputs", {})
    if outputs:
        engine_parts = list(outputs.keys())[:2]  # Max 2 engines in name

    if not engine_parts:
        engine_parts = [result.get("extended_info", {}).get("engine", "analysis")]

    # Extract document title from extended_info
    doc_title = ""
    extended_info = result.get("extended_info", {})
    documents = extended_info.get("documents", [])

    if documents:
        # Get first document title
        first_doc = documents[0]
        doc_title = first_doc.get("title", "") or first_doc.get("id", "")

        # Clean up the title - remove file extension, special chars
        if doc_title:
            # Remove common file extensions
            doc_title = re.sub(r'\.(pdf|txt|md|docx)$', '', doc_title, flags=re.IGNORECASE)
            # Remove author/year patterns like "Author - Year - " at the start
            doc_title = re.sub(r'^[A-Za-z]+\s*-\s*\d{4}\s*-\s*', '', doc_title)
            # Keep only alphanumeric and spaces, replace multiple spaces
            doc_title = re.sub(r'[^a-zA-Z0-9\s]', '', doc_title)
            doc_title = re.sub(r'\s+', '_', doc_title.strip())
            # Truncate to reasonable length (max 50 chars)
            if len(doc_title) > 50:
                doc_title = doc_title[:47] + "..."

    # Build the folder name
    parts = [timestamp]

    # Add engine(s)
    if engine_parts:
        parts.append("_".join(engine_parts[:2]))

    # Add document title if available
    if doc_title:
        parts.append(doc_title)
    else:
        # Fallback to shortened job_id
        parts.append(job_id[:8])

    folder_name = "_".join(parts)

    # Final sanitization - ensure valid folder name
    folder_name = re.sub(r'[<>:"/\\|?*]', '', folder_name)
    folder_name = re.sub(r'_+', '_', folder_name)  # Collapse multiple underscores
    folder_name = folder_name.strip('_')

    return folder_name


# =============================================================================
# MCP TOOLS
# =============================================================================

@mcp.tool()
def get_ai_recommendations(
    document_path: Annotated[str, Field(description="Path to the document to analyze (PDF, TXT, MD, etc.)")],
    max_recommendations: Annotated[int, Field(description="Maximum number of engine recommendations to return (default: 5)")] = 5,
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key for Claude-based recommendations")] = None,
    gemini_api_key: Annotated[Optional[str], Field(description="Gemini API key for recommendations")] = None
) -> str:
    """
    Get AI-powered recommendations for the best analysis engines to use for a document.

    The AI analyzes the document content and suggests 4-5 engines that would provide
    the most valuable insights based on document type, style, and content.

    Returns: JSON with recommended engines, document characteristics, and analysis strategy.
    """
    logger.info(f"Getting AI recommendations for: {document_path}")

    # Read document
    doc = read_document(document_path)
    if "error" in doc:
        return json.dumps({"error": doc["error"]})

    # Extract sample text for curator (up to 2000 chars)
    # Use extracted_text which has actual text content for PDFs
    sample_text = doc.get("extracted_text", doc["content"])[:2000]

    # Build llm_keys for request body
    llm_keys = get_llm_keys(anthropic_api_key, gemini_api_key)

    # Also build headers as backup (for curator endpoint)
    headers = {}
    if llm_keys.get('anthropic_api_key'):
        headers['X-Anthropic-Api-Key'] = llm_keys['anthropic_api_key']
    if llm_keys.get('gemini_api_key'):
        headers['X-Gemini-Api-Key'] = llm_keys['gemini_api_key']

    # Call curator endpoint with llm_keys in body
    result = api_request(
        VISUALIZER_API_URL,
        'POST',
        '/api/analyzer/curator/recommend',
        data={
            "sample_text": sample_text,
            "max_recommendations": max_recommendations,
            "llm_keys": llm_keys  # Include in body
        },
        headers=headers,  # Also in headers as backup
        timeout=120  # Curator may take longer
    )

    if "error" in result:
        return json.dumps({"error": result["error"], "details": result.get("details", "")})

    # Format response for readability
    output = {
        "document": doc["title"],
        "recommendations": result.get("primary_recommendations", []),
        "bundle_recommendations": result.get("bundle_recommendations", []),
        "pipeline_recommendations": result.get("pipeline_recommendations", []),
        "document_characteristics": result.get("document_characteristics", {}),
        "analysis_strategy": result.get("analysis_strategy", ""),
        "curator_note": "The AI has analyzed your document and suggests these engines, bundles, and pipelines for optimal insights."
    }

    num_engines = len(output['recommendations'])
    num_bundles = len(output['bundle_recommendations'])
    num_pipelines = len(output['pipeline_recommendations'])
    logger.info(f"Got {num_engines} engine, {num_bundles} bundle, {num_pipelines} pipeline recommendations")
    return json.dumps(output, indent=2)


@mcp.tool()
def list_available_engines(
    category: Annotated[Optional[str], Field(description="Filter by category (optional)")] = None
) -> str:
    """
    List all available analysis engines.

    Categories include: AI Engines, Argument & Reasoning, Concepts & Frameworks,
    Temporal & Historical, Power & Resources, Evidence & Data, Rhetoric & Language,
    Scholarly Landscape.

    Returns: JSON with engines grouped by category.
    """
    logger.info(f"Listing engines (category: {category or 'all'})")

    result = api_request(VISUALIZER_API_URL, 'GET', '/api/analyzer/engines')

    if "error" in result:
        return json.dumps({"error": result["error"]})

    # Handle if result is a list directly
    engines_list = result if isinstance(result, list) else result.get('engines', [])

    # Filter by category if specified
    if category:
        engines = [e for e in engines_list if e.get('category', '').lower() == category.lower()]
    else:
        engines = engines_list

    # Format for readability - handle missing keys gracefully
    output = {
        "total_engines": len(engines),
        "engines": [
            {
                "key": e.get("key") or e.get("engine_key"),
                "name": e.get("name") or e.get("engine_name"),
                "description": (e.get("description", "") or "")[:100] + "..." if len(e.get("description", "") or "") > 100 else (e.get("description", "") or ""),
                "category": e.get("category")
            }
            for e in engines
        ]
    }

    return json.dumps(output, indent=2)


@mcp.tool()
def submit_analysis(
    document_path: Annotated[str, Field(description="Path to the document to analyze")],
    engine_keys: Annotated[List[str], Field(description="List of engine keys to use (e.g., ['anomaly_detector', 'argument_architecture'])")],
    output_mode: Annotated[str, Field(description="Output mode: 'visual' for 4K images, 'textual' for reports/diagrams")] = "visual",
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key")] = None,
    gemini_api_key: Annotated[Optional[str], Field(description="Gemini API key (required for visual mode)")] = None,
    auto_monitor: Annotated[bool, Field(description="Automatically monitor job and notify when complete (default: True)")] = True
) -> str:
    """
    Submit a document for analysis with selected engines.

    This initiates the analysis pipeline:
    1. Documents are sent to the analyzer backend
    2. Selected engines are executed
    3. Results are generated (visual or textual)
    4. Outputs are uploaded to S3
    5. Job completes and user is notified

    Returns: Job ID for tracking progress.
    """
    logger.info(f"=" * 60)
    logger.info(f"SUBMIT_ANALYSIS: {document_path}")
    logger.info(f"Engines: {engine_keys}")
    logger.info(f"Output mode: {output_mode}, Auto-monitor: {auto_monitor}")
    logger.info(f"=" * 60)

    # Read document
    logger.debug("Reading document...")
    doc = read_document(document_path)
    if "error" in doc:
        logger.error(f"Document read failed: {doc['error']}")
        return json.dumps({"error": doc["error"]})

    logger.info(f"Document loaded: {doc['title']}, size: {doc.get('size', 'unknown')} bytes")

    # Build llm_keys for request body (this is what the visualizer expects!)
    llm_keys = get_llm_keys(anthropic_api_key, gemini_api_key)
    has_anthropic = 'anthropic_api_key' in llm_keys
    has_gemini = 'gemini_api_key' in llm_keys
    logger.debug(f"API keys available - Anthropic: {has_anthropic}, Gemini: {has_gemini}")

    # Validate output mode
    if output_mode not in ['visual', 'textual']:
        logger.error(f"Invalid output_mode: {output_mode}")
        return json.dumps({"error": "output_mode must be 'visual' or 'textual'"})

    # Determine the actual output mode string for the API
    # Valid modes: gemini_image, structured_text_report, executive_memo, mermaid, d3_interactive, etc.
    api_output_mode = "gemini_image" if output_mode == "visual" else "executive_memo"
    logger.debug(f"API output mode: {api_output_mode}")

    # Prepare document for submission (remove extra fields the API doesn't need)
    doc_for_api = {
        "id": doc["id"],
        "title": doc["title"],
        "content": doc["content"],
        "encoding": doc["encoding"]
    }
    logger.debug(f"Document for API: id={doc_for_api['id']}, encoding={doc_for_api['encoding']}, content_size={len(doc_for_api['content']):,} chars")

    job_ids = []
    errors = []

    # Submit each engine separately (more reliable than bundle for arbitrary engine lists)
    for i, engine_key in enumerate(engine_keys, 1):
        logger.info(f"[{i}/{len(engine_keys)}] Submitting engine: {engine_key}")

        result = api_request(
            VISUALIZER_API_URL,
            'POST',
            '/api/analyzer/analyze',
            data={
                "documents": [doc_for_api],
                "engine": engine_key,
                "output_mode": api_output_mode,
                "collection_mode": "single",
                "llm_keys": llm_keys  # Pass API keys in body!
            },
            timeout=120
        )

        if "error" in result:
            error_msg = f"{engine_key}: {result['error']}"
            logger.error(f"[{i}/{len(engine_keys)}] FAILED: {error_msg}")
            if 'details' in result:
                logger.error(f"[{i}/{len(engine_keys)}] Details: {result['details'][:300]}")
            errors.append(error_msg)
        elif result.get("job_id"):
            job_id = result["job_id"]
            logger.info(f"[{i}/{len(engine_keys)}] SUCCESS: job_id={job_id}")
            job_ids.append({
                "engine": engine_key,
                "job_id": job_id
            })
        else:
            error_msg = f"{engine_key}: No job ID returned"
            logger.error(f"[{i}/{len(engine_keys)}] FAILED: {error_msg}")
            logger.error(f"[{i}/{len(engine_keys)}] Full response: {json.dumps(result)[:500]}")
            errors.append(error_msg)

    if not job_ids and errors:
        logger.error(f"ALL SUBMISSIONS FAILED: {len(errors)} errors")
        for err in errors:
            logger.error(f"  - {err}")
        return json.dumps({"error": "All submissions failed", "details": errors})

    logger.info(f"SUBMISSION SUMMARY: {len(job_ids)} succeeded, {len(errors)} failed")

    output = {
        "jobs": job_ids,
        "status": "submitted",
        "document": doc["title"],
        "engines": engine_keys,
        "output_mode": output_mode,
        "errors": errors if errors else None,
        "message": f"Submitted {len(job_ids)} analysis job(s). Use check_job_status() with each job_id to monitor progress."
    }

    # Send notification (silent - no sound for submission)
    send_notification(
        "📊 Analysis Started",
        f"Analyzing {doc['title']} with {len(job_ids)} engine(s)",
        "visualizer,started",
        sound=False
    )

    if auto_monitor and job_ids:
        # Extract just the job IDs and spawn background poller
        all_job_ids = [j["job_id"] for j in job_ids]
        logger.info(f"Spawning background poller for {len(all_job_ids)} jobs")
        spawn_job_poller(all_job_ids)
        output["message"] += " Results will auto-download when complete."
        output["auto_monitor"] = True

    logger.info(f"=" * 60)
    return json.dumps(output, indent=2)


@mcp.tool()
def check_job_status(
    job_id: Annotated[str, Field(description="Job ID returned from submit_analysis()")]
) -> str:
    """
    Check the status of an analysis job.

    Returns: Current job status, progress, and results if complete.
    """
    check_start = time.time()
    job_short = job_id[:8]
    logger.info(f"[{job_short}] Checking job status...")

    result = api_request(
        VISUALIZER_API_URL,
        'GET',
        f'/api/analyzer/jobs/{job_id}'
    )

    if "error" in result:
        logger.error(f"[{job_short}] Failed to get status: {result.get('error')}")
        logger.error(f"[{job_short}] Request details: {result.get('request_id', 'N/A')}, URL: {result.get('url', 'N/A')}")
        return json.dumps({"error": result["error"]})

    status = result.get("status", "unknown")
    created_at = result.get("created_at")
    updated_at = result.get("updated_at")

    # Log timing information
    if created_at and updated_at:
        logger.debug(f"[{job_short}] Created: {created_at}, Updated: {updated_at}")

    # Log detailed status info
    logger.info(f"[{job_short}] Status: {status.upper()}")

    output = {
        "job_id": job_id,
        "status": status,
        "created_at": created_at,
        "updated_at": updated_at
    }

    if status == "completed":
        output["result_available"] = True
        output["message"] = "Job completed! Use get_results() to download outputs."
        logger.info(f"[{job_short}] ✓ Job COMPLETED successfully")
        # Log output info if available
        outputs = result.get("outputs", {})
        if outputs:
            engine_keys = list(outputs.keys())
            logger.info(f"[{job_short}] Engines completed: {engine_keys}")
        send_notification(
            "✅ Job Complete",
            f"Analysis job {job_id[:8]}... is ready",
            "visualizer,completed"
        )
    elif status == "failed":
        error_msg = result.get("error_message") or result.get("error", "Job failed")
        output["error"] = error_msg
        # Log failure details extensively
        logger.error(f"[{job_short}] ✗ Job FAILED: {error_msg}")
        logger.error(f"[{job_short}] Full error response: {json.dumps(result, default=str)[:2000]}")
        # Log any additional debug info from the response
        if "traceback" in result:
            logger.error(f"[{job_short}] Traceback: {result['traceback'][:1000]}")
        if "extended_info" in result:
            ext_info = result["extended_info"]
            logger.error(f"[{job_short}] Engine: {ext_info.get('engine')}, Pipeline: {ext_info.get('pipeline')}")
        send_notification(
            "❌ Job Failed",
            f"Analysis job {job_id[:8]}... failed: {error_msg[:100]}",
            "visualizer,failed"
        )
    elif status in ["pending", "running", "extracting", "curating", "rendering"]:
        output["message"] = f"Job is {status}. Check back in a few moments."
        logger.debug(f"[{job_short}] Job in progress: {status}")
        # Log progress if available
        progress = result.get("progress")
        if progress:
            output["progress"] = progress
            logger.debug(f"[{job_short}] Progress: {progress}")
    else:
        logger.warning(f"[{job_short}] Unknown status: {status}")

    elapsed = time.time() - check_start
    logger.debug(f"[{job_short}] Status check completed in {elapsed:.2f}s")

    return json.dumps(output, indent=2)


@mcp.tool()
def get_results(
    job_id: Annotated[str, Field(description="Job ID to retrieve results for")],
    download_to: Annotated[Optional[str], Field(description="Directory to save results (default: ~/visualizer-results)")] = None
) -> str:
    """
    Retrieve and download completed job results.

    Results are downloaded from S3 and saved to the local filesystem.
    For visual outputs, you'll get 4K PNG images. For textual outputs, you'll get markdown files.

    Returns: Paths to downloaded files.
    """
    download_start = time.time()
    job_short = job_id[:8]
    logger.info(f"[{job_short}] ===== DOWNLOADING RESULTS =====")

    # First get job details
    logger.debug(f"[{job_short}] Fetching job result from API...")
    result = api_request(
        VISUALIZER_API_URL,
        'GET',
        f'/api/analyzer/jobs/{job_id}/result'
    )

    if "error" in result:
        logger.error(f"[{job_short}] Failed to get results: {result.get('error')}")
        logger.error(f"[{job_short}] Request details: {result.get('request_id', 'N/A')}")
        return json.dumps({"error": result["error"]})

    # Log result structure
    logger.debug(f"[{job_short}] Result keys: {list(result.keys())}")
    if "extended_info" in result:
        ext = result["extended_info"]
        logger.info(f"[{job_short}] Engine: {ext.get('engine')}, Documents: {ext.get('documents_total')}")
    if "metadata" in result:
        meta = result["metadata"]
        logger.info(f"[{job_short}] Processing time: {meta.get('total_ms')}ms, Cost: ${meta.get('cost_usd', 0):.4f}")

    # Determine output directory with meaningful folder name
    output_dir = Path(download_to).expanduser() if download_to else OUTPUT_DIR
    folder_name = generate_meaningful_folder_name(result, job_id)
    job_dir = output_dir / folder_name
    job_dir.mkdir(parents=True, exist_ok=True)

    logger.info(f"[{job_short}] Output folder: {folder_name}")
    logger.debug(f"[{job_short}] Full path: {job_dir}")

    downloaded_files = []
    download_errors = []
    total_bytes = 0

    # Handle the CORRECT response structure: outputs.{engine_key}.{image_url|text|...}
    outputs = result.get("outputs", {})
    logger.info(f"[{job_short}] Found {len(outputs)} output(s) to download")

    if isinstance(outputs, dict):
        for engine_key, engine_result in outputs.items():
            if not isinstance(engine_result, dict):
                logger.warning(f"[{job_short}] Skipping {engine_key}: not a dict ({type(engine_result)})")
                continue

            logger.debug(f"[{job_short}] Processing output: {engine_key}")

            # Check for image_url (visual mode - Gemini image)
            image_url = engine_result.get("image_url")
            if image_url:
                # Determine file extension from URL or default to .jpg
                if ".png" in image_url.lower():
                    ext = ".png"
                elif ".jpg" in image_url.lower() or ".jpeg" in image_url.lower():
                    ext = ".jpg"
                else:
                    ext = ".jpg"  # Default for S3 images

                file_path = job_dir / f"{engine_key}{ext}"
                logger.debug(f"[{job_short}] Downloading image: {engine_key}{ext}")
                dl_start = time.time()
                success, error = download_file_from_url(image_url, file_path)
                dl_elapsed = time.time() - dl_start
                if success:
                    file_size = file_path.stat().st_size
                    total_bytes += file_size
                    logger.info(f"[{job_short}] ✓ {engine_key}{ext} ({file_size/1024:.1f}KB in {dl_elapsed:.1f}s)")
                    downloaded_files.append(str(file_path))
                else:
                    logger.error(f"[{job_short}] ✗ Failed {engine_key}: {error}")
                    download_errors.append(f"{engine_key}: {error}")
                continue

            # Check for text content (textual mode)
            text_content = engine_result.get("text") or engine_result.get("content") or engine_result.get("output")
            if text_content and isinstance(text_content, str):
                file_path = job_dir / f"{engine_key}.md"
                file_path.write_text(text_content, encoding='utf-8')
                file_size = len(text_content.encode('utf-8'))
                total_bytes += file_size
                downloaded_files.append(str(file_path))
                logger.info(f"[{job_short}] ✓ {engine_key}.md ({file_size/1024:.1f}KB)")
                continue

            # Check for S3 URL in output field (legacy format)
            output_data = engine_result.get("output")
            if isinstance(output_data, str):
                if output_data.startswith("s3://"):
                    # Convert S3 URL to HTTP URL
                    parts = output_data[5:].split("/", 1)
                    if len(parts) == 2:
                        bucket, key = parts
                        http_url = f"https://{bucket}.s3.amazonaws.com/{key}"
                        ext = ".png" if any(x in key.lower() for x in ['.png', '.jpg', '.jpeg', 'image']) else ".md"
                        file_path = job_dir / f"{engine_key}{ext}"
                        success, error = download_file_from_url(http_url, file_path)
                        if success:
                            downloaded_files.append(str(file_path))
                        else:
                            download_errors.append(f"{engine_key}: {error}")
                elif output_data.startswith("http"):
                    ext = ".png" if any(x in output_data.lower() for x in ['.png', '.jpg', '.jpeg']) else ".md"
                    file_path = job_dir / f"{engine_key}{ext}"
                    success, error = download_file_from_url(output_data, file_path)
                    if success:
                        downloaded_files.append(str(file_path))
                    else:
                        download_errors.append(f"{engine_key}: {error}")
                else:
                    # Inline text content
                    file_path = job_dir / f"{engine_key}.md"
                    file_path.write_text(output_data, encoding='utf-8')
                    downloaded_files.append(str(file_path))

    # Also check legacy "result" structure for backwards compatibility
    legacy_results = result.get("result", {})
    if isinstance(legacy_results, dict) and not outputs:
        for engine_key, engine_result in legacy_results.items():
            if isinstance(engine_result, dict) and "output" in engine_result:
                output_data = engine_result["output"]
                if isinstance(output_data, str):
                    if output_data.startswith("http"):
                        ext = ".png" if "image" in output_data.lower() else ".md"
                        file_path = job_dir / f"{engine_key}{ext}"
                        success, error = download_file_from_url(output_data, file_path)
                        if success:
                            downloaded_files.append(str(file_path))
                        else:
                            download_errors.append(f"{engine_key}: {error}")
                    else:
                        file_path = job_dir / f"{engine_key}.md"
                        file_path.write_text(output_data, encoding='utf-8')
                        downloaded_files.append(str(file_path))

    # Save raw JSON for reference
    json_path = job_dir / "results.json"
    json_content = json.dumps(result, indent=2)
    json_path.write_text(json_content, encoding='utf-8')
    json_size = len(json_content.encode('utf-8'))
    total_bytes += json_size
    downloaded_files.append(str(json_path))
    logger.debug(f"[{job_short}] Saved results.json ({json_size/1024:.1f}KB)")

    # Calculate totals
    total_elapsed = time.time() - download_start
    total_kb = total_bytes / 1024
    total_mb = total_bytes / (1024 * 1024)

    output = {
        "job_id": job_id,
        "download_directory": str(job_dir),
        "files_downloaded": len(downloaded_files),
        "files": downloaded_files,
        "total_size_kb": round(total_kb, 1),
        "download_time_seconds": round(total_elapsed, 1),
        "errors": download_errors if download_errors else None
    }

    # Log summary
    logger.info(f"[{job_short}] ===== DOWNLOAD COMPLETE =====")
    logger.info(f"[{job_short}] Files: {len(downloaded_files)}, Total: {total_mb:.2f}MB, Time: {total_elapsed:.1f}s")
    if download_errors:
        logger.warning(f"[{job_short}] Errors: {len(download_errors)} - {download_errors}")
    logger.info(f"[{job_short}] Location: {job_dir}")

    # Send notification
    send_notification(
        "✅ Results Downloaded",
        f"Downloaded {len(downloaded_files)} file(s) ({total_mb:.1f}MB) to {folder_name}",
        "visualizer,completed"
    )

    return json.dumps(output, indent=2)


@mcp.tool()
def list_bundles() -> str:
    """
    List available engine bundles.

    Bundles are pre-configured groups of engines that work well together.

    Returns: JSON with available bundles.
    """
    result = api_request(VISUALIZER_API_URL, 'GET', '/api/analyzer/bundles')

    if "error" in result:
        return json.dumps({"error": result["error"]})

    return json.dumps({"bundles": result}, indent=2)


@mcp.tool()
def list_pipelines() -> str:
    """
    List available analysis pipelines.

    Pipelines are sequences of engines where the output of one feeds into the next.

    Returns: JSON with available pipelines.
    """
    result = api_request(VISUALIZER_API_URL, 'GET', '/api/analyzer/pipelines')

    if "error" in result:
        return json.dumps({"error": result["error"]})

    return json.dumps({"pipelines": result}, indent=2)


# =============================================================================
# BATCH ANALYSIS TOOLS
# =============================================================================

@mcp.tool()
def scan_folder(
    folder_path: Annotated[str, Field(description="Path to folder containing documents to analyze")],
    file_types: Annotated[str, Field(description="Comma-separated file extensions to include (default: 'pdf,txt,md')")] = "pdf,txt,md"
) -> str:
    """
    Scan a folder and list all documents available for batch analysis.

    Use this to preview what files would be processed before submitting a batch.

    Returns: JSON with list of files found, sizes, and types.
    """
    logger.info(f"Scanning folder: {folder_path}")

    path = Path(folder_path).expanduser().resolve()

    if not path.exists():
        return json.dumps({"error": f"Folder not found: {folder_path}"})

    if not path.is_dir():
        return json.dumps({"error": f"Not a directory: {folder_path}"})

    # Parse file types
    extensions = {f".{ext.strip().lower().lstrip('.')}" for ext in file_types.split(",")}

    files = []
    total_size = 0

    for file_path in path.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in extensions:
            size = file_path.stat().st_size
            files.append({
                "name": file_path.name,
                "path": str(file_path),
                "type": file_path.suffix.lower(),
                "size_bytes": size,
                "size_human": f"{size / 1024:.1f} KB" if size < 1024 * 1024 else f"{size / (1024 * 1024):.1f} MB"
            })
            total_size += size

    # Sort by name
    files.sort(key=lambda x: x["name"])

    output = {
        "folder": str(path),
        "file_types_searched": list(extensions),
        "files_found": len(files),
        "total_size_bytes": total_size,
        "total_size_human": f"{total_size / 1024:.1f} KB" if total_size < 1024 * 1024 else f"{total_size / (1024 * 1024):.1f} MB",
        "files": files
    }

    return json.dumps(output, indent=2)


@mcp.tool()
def submit_batch_analysis(
    folder_path: Annotated[str, Field(description="Path to folder containing documents")],
    engine_keys: Annotated[List[str], Field(description="List of engine keys to run on EACH document")],
    output_mode: Annotated[str, Field(description="Output mode: 'visual' for 4K images, 'textual' for reports")] = "visual",
    file_types: Annotated[str, Field(description="Comma-separated extensions to process (default: 'pdf,txt,md')")] = "pdf,txt,md",
    max_documents: Annotated[int, Field(description="Maximum number of documents to process (default: 20)")] = 20,
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key")] = None,
    gemini_api_key: Annotated[Optional[str], Field(description="Gemini API key (required for visual mode)")] = None
) -> str:
    """
    Submit an entire folder of documents for batch analysis.

    Each document will be analyzed with ALL selected engines, generating N × M jobs
    (N documents × M engines).

    Returns: Grouped job IDs for tracking each document's analysis.
    """
    logger.info(f"Batch analysis: {folder_path} with {len(engine_keys)} engines")

    # First scan the folder
    scan_result = json.loads(scan_folder(folder_path, file_types))
    if "error" in scan_result:
        return json.dumps(scan_result)

    files = scan_result.get("files", [])
    if not files:
        return json.dumps({"error": f"No matching files found in {folder_path}"})

    # Limit number of documents
    if len(files) > max_documents:
        logger.warning(f"Limiting to {max_documents} documents (found {len(files)})")
        files = files[:max_documents]

    # Build llm_keys
    llm_keys = get_llm_keys(anthropic_api_key, gemini_api_key)

    # Determine API output mode
    api_output_mode = "gemini_image" if output_mode == "visual" else "executive_memo"

    batch_results = []
    total_jobs = 0
    total_errors = 0

    for file_info in files:
        file_path = file_info["path"]
        doc = read_document(file_path)

        if "error" in doc:
            batch_results.append({
                "file": file_info["name"],
                "status": "error",
                "error": doc["error"],
                "jobs": []
            })
            total_errors += 1
            continue

        doc_for_api = {
            "id": doc["id"],
            "title": doc["title"],
            "content": doc["content"],
            "encoding": doc["encoding"]
        }

        file_jobs = []
        file_errors = []

        for engine_key in engine_keys:
            result = api_request(
                VISUALIZER_API_URL,
                'POST',
                '/api/analyzer/analyze',
                data={
                    "documents": [doc_for_api],
                    "engine": engine_key,
                    "output_mode": api_output_mode,
                    "collection_mode": "single",
                    "llm_keys": llm_keys
                },
                timeout=120
            )

            if "error" in result:
                file_errors.append(f"{engine_key}: {result['error']}")
            elif result.get("job_id"):
                file_jobs.append({
                    "engine": engine_key,
                    "job_id": result["job_id"]
                })
                total_jobs += 1
            else:
                file_errors.append(f"{engine_key}: No job ID returned")

        batch_results.append({
            "file": file_info["name"],
            "status": "submitted" if file_jobs else "failed",
            "jobs": file_jobs,
            "errors": file_errors if file_errors else None
        })

    # Send notification (silent - no sound for submission)
    send_notification(
        "📊 Batch Analysis Started",
        f"Processing {len(files)} documents with {len(engine_keys)} engines ({total_jobs} total jobs)",
        "visualizer,batch,started",
        sound=False
    )

    # Collect all job IDs for auto-download
    all_job_ids = []
    for doc_result in batch_results:
        for job_info in doc_result.get("jobs", []):
            if job_info.get("job_id"):
                all_job_ids.append(job_info["job_id"])

    # Spawn background poller for auto-download
    if all_job_ids:
        spawn_job_poller(all_job_ids)

    output = {
        "batch_status": "submitted",
        "folder": folder_path,
        "documents_processed": len(files),
        "engines_per_document": len(engine_keys),
        "total_jobs": total_jobs,
        "total_errors": total_errors,
        "output_mode": output_mode,
        "documents": batch_results,
        "message": f"Submitted {total_jobs} job(s) for {len(files)} document(s). Results will auto-download when complete.",
        "auto_monitor": True if all_job_ids else False
    }

    logger.info(f"Batch submitted: {total_jobs} jobs for {len(files)} documents")
    return json.dumps(output, indent=2)


# =============================================================================
# PIPELINE ANALYSIS TOOLS
# =============================================================================

@mcp.tool()
def submit_pipeline_analysis(
    document_path: Annotated[str, Field(description="Path to the document to analyze")],
    pipeline_key: Annotated[str, Field(description="Pipeline key (e.g., 'argument_deep_dive', 'evidence_chain')")],
    output_mode: Annotated[str, Field(description="Output mode: 'visual' for 4K images, 'textual' for reports")] = "visual",
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key")] = None,
    gemini_api_key: Annotated[Optional[str], Field(description="Gemini API key (required for visual mode)")] = None
) -> str:
    """
    Submit a document for pipeline analysis.

    Pipelines chain multiple engines sequentially, where each engine's output
    feeds into the next. This provides deeper, multi-stage analysis.

    Returns: Job ID for the pipeline execution.
    """
    logger.info(f"Pipeline analysis: {document_path} with pipeline '{pipeline_key}'")

    # Read document
    doc = read_document(document_path)
    if "error" in doc:
        return json.dumps({"error": doc["error"]})

    # Build llm_keys
    llm_keys = get_llm_keys(anthropic_api_key, gemini_api_key)

    # Determine API output mode
    api_output_mode = "gemini_image" if output_mode == "visual" else "executive_memo"

    # Prepare document
    doc_for_api = {
        "id": doc["id"],
        "title": doc["title"],
        "content": doc["content"],
        "encoding": doc["encoding"]
    }

    # Submit pipeline analysis (use /analyze/pipeline endpoint)
    result = api_request(
        VISUALIZER_API_URL,
        'POST',
        '/api/analyzer/analyze/pipeline',
        data={
            "documents": [doc_for_api],
            "pipeline": pipeline_key,
            "output_mode": api_output_mode,
            "include_intermediate_outputs": True,
            "llm_keys": llm_keys
        },
        timeout=120
    )

    if "error" in result:
        return json.dumps({"error": result["error"], "details": result.get("details", "")})

    job_id = result.get("job_id")
    if not job_id:
        return json.dumps({"error": "No job ID returned from pipeline submission"})

    output = {
        "job_id": job_id,
        "status": "submitted",
        "document": doc["title"],
        "pipeline": pipeline_key,
        "output_mode": output_mode,
        "message": f"Pipeline '{pipeline_key}' started. Use check_job_status('{job_id}') to monitor progress."
    }

    # Send notification (silent - no sound for submission)
    send_notification(
        "🔗 Pipeline Analysis Started",
        f"Running pipeline '{pipeline_key}' on {doc['title']}",
        "visualizer,pipeline,started",
        sound=False
    )

    # Spawn background poller for auto-download
    spawn_job_poller([job_id])
    output["message"] += " Results will auto-download when complete."
    output["auto_monitor"] = True

    logger.info(f"Pipeline submitted: {job_id}")
    return json.dumps(output, indent=2)


# =============================================================================
# Phase 6: Intent-Based Analysis Tools
# =============================================================================

@mcp.tool()
def get_collection_affordances(
    document_paths: Annotated[List[str], Field(description="List of document paths to analyze for affordances")],
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key")] = None
) -> str:
    """
    Detect what types of analysis a document collection supports.

    This is a diagnostic tool that samples your documents and determines:
    - Domain (policy, finance, philosophy, technology, etc.)
    - Entity density (high/medium/low)
    - Temporal content (high/medium/low)
    - Quantitative content (high/medium/low)
    - Which engine categories are suitable vs unsuitable

    Use this to understand what analyses make sense for your collection
    before running analyze_collection_with_intent().

    Returns: JSON with affordances and suitable/unsuitable engine categories.
    """
    # Build document list from paths
    documents = []
    for path in document_paths[:20]:  # Limit to 20 docs
        doc = prepare_document(path)
        if doc:
            documents.append({
                "id": str(Path(path).stem),
                "title": doc.get("title", Path(path).stem),
                "content": doc.get("content", "")[:2000],  # First 2000 chars
                "source": doc.get("source"),
                "date": doc.get("date"),
            })

    if not documents:
        return json.dumps({"error": "No valid documents found"}, indent=2)

    # Build headers with API keys
    llm_keys = get_llm_keys(anthropic_api_key=anthropic_api_key)
    headers = build_llm_headers(llm_keys)

    # Call the affordances endpoint
    try:
        response = api_request(
            ANALYZER_API_URL,
            'POST',
            '/v1/curator/affordances',
            data={
                "documents": documents,
                "sample_chars_per_doc": 500,
                "max_docs_to_sample": 10,
            },
            headers=headers,
            timeout=60,
        )

        logger.info(f"Affordances detected for {len(documents)} documents")
        return json.dumps(response, indent=2)

    except Exception as e:
        return json.dumps({
            "error": f"Failed to detect affordances: {str(e)}",
            "documents_submitted": len(documents),
        }, indent=2)


@mcp.tool()
def classify_intent(
    user_request: Annotated[str, Field(description="Natural language description of what you want to understand")],
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key")] = None
) -> str:
    """
    Classify a natural language intent into verb+noun taxonomy.

    Instead of specifying engine names, describe what you want to understand:
    - "Map the key players and their relationships"
    - "Trace how this concept evolved over time"
    - "Compare approaches across different jurisdictions"
    - "Find gaps in the current research"
    - "Track the money flows"

    Returns: JSON with primary_verb, primary_noun, and confidence.
    """
    # Build headers with API keys
    llm_keys = get_llm_keys(anthropic_api_key=anthropic_api_key)
    headers = build_llm_headers(llm_keys)

    try:
        response = api_request(
            ANALYZER_API_URL,
            'POST',
            '/v1/curator/classify-intent',
            data={"user_request": user_request},
            headers=headers,
            timeout=30,
        )

        logger.info(f"Intent classified: {response.get('primary_verb')} + {response.get('primary_noun')}")
        return json.dumps(response, indent=2)

    except Exception as e:
        return json.dumps({
            "error": f"Failed to classify intent: {str(e)}",
        }, indent=2)


@mcp.tool()
def analyze_collection_with_intent(
    document_paths: Annotated[List[str], Field(description="List of document paths to analyze")],
    intent: Annotated[str, Field(description="What you want to understand (e.g., 'Map the key players')")],
    output_modes: Annotated[Optional[List[str]], Field(description="Output formats: ['gemini_image', 'smart_table', 'text']. Default: all three.")] = None,
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key")] = None,
    gemini_api_key: Annotated[Optional[str], Field(description="Gemini API key for visual output")] = None
) -> str:
    """
    Analyze a collection of documents based on what you want to understand.

    Instead of selecting engines, describe your intent:
    - "Map the key players and their relationships"
    - "Trace how this concept evolved over time"
    - "Compare approaches across different jurisdictions"
    - "Find gaps in the current research"
    - "Track the money flows"
    - "Evaluate the strength of the arguments"
    - "Synthesize the main themes"

    The AI will:
    1. Sample the collection to detect what analyses it supports
    2. Classify your intent into verb+noun taxonomy
    3. Select the best engine for your intent
    4. Extract from each document, synthesize across collection
    5. Generate outputs in multiple formats (image, table, text)

    Returns: Job ID for tracking. Results auto-download when complete.
    """
    # Default output modes
    if output_modes is None:
        output_modes = ["gemini_image", "smart_table", "text"]

    # Build headers with API keys
    llm_keys = get_llm_keys(anthropic_api_key=anthropic_api_key, gemini_api_key=gemini_api_key)
    headers = build_llm_headers(llm_keys)

    output = {
        "intent": intent,
        "document_count": len(document_paths),
        "output_modes": output_modes,
    }

    # Step 1: Prepare documents
    documents = []
    for path in document_paths:
        doc = prepare_document(path)
        if doc:
            documents.append(doc)

    if not documents:
        return json.dumps({"error": "No valid documents found"}, indent=2)

    output["prepared_documents"] = len(documents)

    # Step 2: Classify intent
    try:
        intent_response = api_request(
            ANALYZER_API_URL,
            'POST',
            '/v1/curator/classify-intent',
            data={"user_request": intent},
            headers=headers,
            timeout=30,
        )
        output["classified_intent"] = {
            "verb": intent_response.get("primary_verb"),
            "noun": intent_response.get("primary_noun"),
            "confidence": intent_response.get("confidence"),
        }
        logger.info(f"Intent: {intent_response.get('primary_verb')} + {intent_response.get('primary_noun')}")
    except Exception as e:
        logger.warning(f"Intent classification failed: {e}")
        output["classified_intent"] = {"error": str(e)}
        # Continue without intent classification

    # Step 3: Get AI recommendations with intent
    sample_text = "\n\n---\n\n".join([
        f"[{d['title']}]\n{d['content'][:500]}"
        for d in documents[:5]
    ])

    try:
        recommend_data = {
            "sample_text": sample_text,
            "analysis_goal": intent,
            "max_recommendations": 3,
        }

        # Add intent if we have it
        if output.get("classified_intent") and not output["classified_intent"].get("error"):
            recommend_data["intent"] = {
                "verb": output["classified_intent"]["verb"],
                "noun": output["classified_intent"]["noun"],
            }

        recommend_response = api_request(
            ANALYZER_API_URL,
            'POST',
            '/v1/curator/recommend',
            data=recommend_data,
            headers=headers,
            timeout=60,
        )

        # Get top engine recommendation
        recommendations = recommend_response.get("primary_recommendations", [])
        if not recommendations:
            return json.dumps({"error": "No engine recommendations returned"}, indent=2)

        top_engine = recommendations[0]
        engine_key = top_engine.get("engine_key")
        output["selected_engine"] = {
            "engine_key": engine_key,
            "engine_name": top_engine.get("engine_name"),
            "confidence": top_engine.get("confidence"),
            "rationale": top_engine.get("rationale"),
            "recommended_outputs": top_engine.get("recommended_outputs", []),
        }
        logger.info(f"Selected engine: {engine_key}")

    except Exception as e:
        return json.dumps({
            "error": f"Engine recommendation failed: {str(e)}",
            **output,
        }, indent=2)

    # Step 4: Submit analysis with multi-output
    primary_output_mode = output_modes[0] if output_modes else "gemini_image"

    try:
        submit_response = api_request(
            ANALYZER_API_URL,
            'POST',
            '/v1/analyze',
            data={
                "documents": documents,
                "engine": engine_key,
                "output_mode": primary_output_mode,
                "output_modes": output_modes,  # Phase 4: Multi-output
            },
            headers=headers,
            timeout=30,
        )

        job_id = submit_response.get("job_id")
        output["job_id"] = job_id
        output["status"] = "submitted"
        output["message"] = f"Analysis job submitted with {len(output_modes)} output modes"

        # Spawn background poller for auto-download
        spawn_job_poller([job_id])
        output["message"] += ". Results will auto-download when complete."
        output["auto_monitor"] = True

        logger.info(f"Intent-based analysis submitted: {job_id}")

    except Exception as e:
        return json.dumps({
            "error": f"Analysis submission failed: {str(e)}",
            **output,
        }, indent=2)

    return json.dumps(output, indent=2)


@mcp.tool()
def analyze_folder_with_intent(
    folder_path: Annotated[str, Field(description="Path to folder containing documents")],
    intent: Annotated[str, Field(description="What you want to understand")],
    file_types: Annotated[str, Field(description="Comma-separated extensions (default: 'pdf,txt,md')")] = "pdf,txt,md",
    max_documents: Annotated[int, Field(description="Maximum documents to process (default: 20)")] = 20,
    output_modes: Annotated[Optional[List[str]], Field(description="Output formats")] = None,
    anthropic_api_key: Annotated[Optional[str], Field(description="Anthropic API key")] = None,
    gemini_api_key: Annotated[Optional[str], Field(description="Gemini API key")] = None
) -> str:
    """
    Analyze a folder of documents with natural language intent.

    Convenience wrapper around analyze_collection_with_intent() that
    scans a folder for documents first.

    Returns: Job ID for tracking.
    """
    # Scan folder for documents
    folder = Path(folder_path).expanduser().resolve()
    if not folder.exists():
        return json.dumps({"error": f"Folder not found: {folder_path}"}, indent=2)

    extensions = set(f".{ext.strip().lower()}" for ext in file_types.split(","))

    document_paths = []
    for ext in extensions:
        document_paths.extend(folder.glob(f"*{ext}"))

    document_paths = sorted(document_paths)[:max_documents]

    if not document_paths:
        return json.dumps({
            "error": f"No documents found in {folder_path} with extensions {file_types}",
        }, indent=2)

    # Call the main intent analysis function
    return analyze_collection_with_intent(
        document_paths=[str(p) for p in document_paths],
        intent=intent,
        output_modes=output_modes,
        anthropic_api_key=anthropic_api_key,
        gemini_api_key=gemini_api_key,
    )


if __name__ == "__main__":
    # Run the MCP server
    mcp.run()
